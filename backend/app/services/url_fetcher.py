"""Fetch and extract text from a URL (HTML, PDF, DOCX, XLSX), with a Playwright fallback
for thin / JS-rendered / blocked pages."""
from __future__ import annotations

import asyncio
import hashlib
import logging
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from io import BytesIO

import httpx
from bs4 import BeautifulSoup

from app.core.config import settings

logger = logging.getLogger(__name__)

# Declared bot UA with contact — required by SEC.gov's fair-access policy (a generic
# browser UA gets 403 from SEC). Browser-like spoofing helps a few sites but breaks more.
# Public so the robots.txt gate matches its product token against ``User-agent:`` lines.
USER_AGENT = "RFPIntelligenceBot/1.0 (+mailto:admin@example.com)"
_USER_AGENT = USER_AGENT  # internal alias kept for existing references
# Browser-like headers used only as a retry when the declared bot UA is refused
# (401/403/429). Some sites gate non-browser UAs; we try this before a full JS render.
_BROWSER_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}
_MAX_XLSX_ROWS = 2000

# Lazily-launched, process-wide headless browser (reused across renders).
_playwright = None
_browser = None


@dataclass
class FetchedContent:
    """Result of fetching and extracting content from a URL."""

    url: str
    title: str
    text: str
    content_type: str  # "html" | "pdf" | "docx" | "xlsx"
    content_hash: str
    content_length: int
    raw_html: str | None = None  # set for HTML pages, for link/feed extraction


def _extract_text_from_html(raw_html: str) -> tuple[str, str]:
    """Return (title, cleaned text) from raw HTML.

    Scripts/styles are always dropped. Layout "chrome" (nav/header/footer/aside/form)
    is dropped only when real content survives: ASP.NET WebForms pages wrap their entire
    body in a single ``<form runat="server">``, so blanket form-removal would discard the
    whole page — which is what was silently breaking RBI and other ``.aspx`` government /
    bank sites (the fetch returned 200, but extraction yielded nothing).
    """
    soup = BeautifulSoup(raw_html, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    title = soup.title.get_text(strip=True) if soup.title else ""

    def _clean(node) -> str:
        raw = node.get_text(separator="\n", strip=True)
        return "\n".join(line.strip() for line in raw.splitlines() if line.strip())

    main = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", {"role": "main"})
        or soup.find("div", id="content")
        or soup.find("div", class_="content")
    )
    main_text = _clean(main) if main else ""

    # Body text before vs. after stripping layout chrome. If stripping removes most of
    # the text, the "chrome" was actually the content (single-<form> WebForms wrapper) —
    # keep the un-stripped body rather than returning an empty string.
    body_full = _clean(soup.body or soup)
    for tag in soup(["nav", "footer", "header", "aside", "form"]):
        tag.decompose()
    body_stripped = _clean(soup.body or soup)
    body_text = body_stripped if len(body_stripped) >= len(body_full) * 0.5 else body_full

    text = main_text
    if len(text) < 200 and len(body_text) > len(text):
        text = body_text
    return title, text


def _extract_text_from_pdf(pdf_bytes: bytes) -> tuple[str, str]:
    """Return (title, extracted text) from PDF bytes."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise RuntimeError("PyPDF2 is required for PDF extraction.")

    reader = PdfReader(BytesIO(pdf_bytes))
    pages_text = [page.extract_text() for page in reader.pages if page.extract_text()]
    full_text = "\n".join(pages_text)

    title = ""
    if reader.metadata and reader.metadata.title:
        title = reader.metadata.title
    if not title:
        for line in full_text.splitlines():
            if line.strip():
                title = line.strip()[:200]
                break
    return title, full_text


def _extract_text_from_docx(data: bytes) -> tuple[str, str]:
    """Return (title, text) from .docx bytes (paragraphs + table cells)."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX extraction.")

    doc = DocxDocument(BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    title = parts[0][:200] if parts else "Document"
    return title, text


def _extract_text_from_xlsx(data: bytes) -> tuple[str, str]:
    """Return (title, text) from .xlsx bytes (sheets -> rows -> tab-joined cells)."""
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    try:
        for sheet in workbook.worksheets:
            lines.append(f"# {sheet.title}")
            for index, row in enumerate(sheet.iter_rows(values_only=True)):
                if index >= _MAX_XLSX_ROWS:
                    break
                cells = [str(cell) for cell in row if cell is not None]
                if cells:
                    lines.append(" | ".join(cells))
    finally:
        workbook.close()
    return "Spreadsheet", "\n".join(lines)


_SOFFICE_BIN: str | None = None  # cached LibreOffice path ("" = looked up, not found)


def _find_soffice() -> str | None:
    """Locate the LibreOffice ``soffice`` binary (PATH or common install dirs); cache result."""
    global _SOFFICE_BIN
    if _SOFFICE_BIN is not None:
        return _SOFFICE_BIN or None
    for name in ("soffice", "soffice.exe", "libreoffice"):
        found = shutil.which(name)
        if found:
            _SOFFICE_BIN = found
            return found
    for candidate in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/opt/libreoffice/program/soffice",
    ):
        if os.path.exists(candidate):
            _SOFFICE_BIN = candidate
            return candidate
    _SOFFICE_BIN = ""  # cache the negative lookup
    return None


async def _extract_text_from_doc(data: bytes) -> tuple[str, str]:
    """Return (title, text) from a legacy binary ``.doc`` / ``.rtf`` via LibreOffice headless.

    Requires LibreOffice (``soffice``) on the host. Raises a clear ``RuntimeError`` when it is
    absent or conversion fails, so the caller can skip the file instead of crashing the crawl.
    """
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not found; cannot parse legacy .doc/.rtf files. "
            "Install LibreOffice or disable these sources."
        )

    def _convert() -> str:
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "input.doc")
            with open(src, "wb") as handle:
                handle.write(data)
            subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmp, src],
                check=True, capture_output=True, timeout=120,
            )
            out = os.path.join(tmp, "input.txt")
            if not os.path.exists(out):
                raise RuntimeError("LibreOffice produced no text output for the .doc file.")
            with open(out, "r", encoding="utf-8", errors="replace") as handle:
                return handle.read()

    text = await asyncio.to_thread(_convert)
    title = next((line.strip()[:200] for line in text.splitlines() if line.strip()), "")
    return title or "Document", text


async def _get_browser():
    """Lazily launch a process-wide headless Chromium (reused across renders)."""
    global _playwright, _browser
    if _browser is not None:
        return _browser
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        logger.info("Playwright not installed; skipping JS rendering.")
        return None
    try:
        _playwright = await async_playwright().start()
        _browser = await _playwright.chromium.launch(headless=True)
        return _browser
    except Exception as exc:  # noqa: BLE001 - degrade to httpx-only
        logger.warning("Playwright launch failed (JS rendering disabled): %s", exc)
        _playwright = None
        _browser = None
        return None


_LOAD_MORE_TOKENS = ("load more", "show more", "view more", "more results", "load older", "see more")


async def _drive_interactions(page) -> list[str]:
    """Click "Load More" controls, walk ``__doPostBack`` pagers, and trigger infinite scroll.

    Load-more / scroll content accumulates in the live DOM (captured by the final
    ``page.content()``); post-back pager clicks replace the DOM, so each such page is captured
    as a separate snapshot and returned for the caller to concatenate. Best-effort throughout.
    """
    snapshots: list[str] = []
    prev_height = 0
    stagnant = 0
    seen_pages: set[str] = set()
    for _ in range(max(1, settings.CRAWL_INTERACT_MAX_CLICKS)):
        acted = False
        # 1) "Load more" / "show more" buttons or links.
        for token in _LOAD_MORE_TOKENS:
            try:
                loc = page.locator(
                    f"button:has-text('{token}'), a:has-text('{token}'), [role=button]:has-text('{token}')"
                )
                if await loc.count() and await loc.first.is_visible():
                    await loc.first.click(timeout=3000)
                    await page.wait_for_timeout(1200)
                    acted = True
                    break
            except Exception:  # noqa: BLE001 - control may detach mid-interaction
                continue
        # 2) ASP.NET __doPostBack numbered pager (GridView/Repeater): advance one page.
        if not acted:
            try:
                pagers = page.locator("a[href*='__doPostBack']")
                count = min(await pagers.count(), 12)
                for index in range(count):
                    label = (await pagers.nth(index).inner_text() or "").strip()
                    if label.isdigit() and label not in seen_pages and label != "1":
                        seen_pages.add(label)
                        await pagers.nth(index).click(timeout=3000)
                        await page.wait_for_timeout(1300)
                        snapshots.append(await page.content())
                        acted = True
                        break
            except Exception:  # noqa: BLE001
                pass
        # 3) Infinite scroll.
        try:
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(900)
            height = int(await page.evaluate("document.body.scrollHeight"))
        except Exception:  # noqa: BLE001
            height = prev_height
        if acted or height > prev_height:
            stagnant = 0
        else:
            stagnant += 1
            if stagnant >= 2:  # two quiet rounds -> nothing left to load
                break
        prev_height = height
    return snapshots


async def _render_with_playwright(url: str, timeout: float = 30.0, interact: bool = False) -> str | None:
    """Render a page with a real browser and return its HTML, or None on failure.

    When ``interact`` is set, also drive Load-More/infinite-scroll/pager interactions so
    JS-populated listings are fully expanded before the HTML is captured.
    """
    browser = await _get_browser()
    if browser is None:
        return None
    page = None
    try:
        page = await browser.new_page()
        await page.goto(url, wait_until="domcontentloaded", timeout=int(timeout * 1000))
        try:
            await page.wait_for_load_state("networkidle", timeout=5000)
        except Exception:  # noqa: BLE001 - networkidle is best-effort
            pass
        if interact and settings.CRAWL_INTERACT_JS:
            snapshots = await _drive_interactions(page)
            final = await page.content()
            return "\n".join([*snapshots, final]) if snapshots else final
        return await page.content()
    except Exception as exc:  # noqa: BLE001
        logger.warning("Playwright render failed for %s: %s", url, exc)
        return None
    finally:
        if page is not None:
            try:
                await page.close()
            except Exception:  # noqa: BLE001
                pass


async def render_interactive(url: str, timeout: float = 30.0) -> str | None:
    """Public: render a page with JS + Load-More/scroll/pager interaction. Returns HTML or None.

    Used by the ingestion layer to re-expand a listing that yielded too few static candidates.
    """
    if not (settings.CRAWL_RENDER_JS and settings.CRAWL_INTERACT_JS):
        return None
    return await _render_with_playwright(url, timeout, interact=True)


async def shutdown_browser() -> None:
    """Close the shared browser (call on app shutdown if desired)."""
    global _playwright, _browser
    try:
        if _browser is not None:
            await _browser.close()
        if _playwright is not None:
            await _playwright.stop()
    except Exception:  # noqa: BLE001
        pass
    finally:
        _browser = None
        _playwright = None


# Transient server/network failures worth a quick retry — RBI and similar government
# portals intermittently return 502/503/504 under load.
_TRANSIENT_STATUS = (502, 503, 504, 429)
_MAX_RETRIES = 2  # extra attempts after the first (3 total)


async def _http_get(url: str, timeout: float, headers: dict | None = None) -> httpx.Response:
    headers = headers or {
        "User-Agent": _USER_AGENT,
        "Accept": "text/html,application/pdf,application/rss+xml,*/*",
    }
    async with httpx.AsyncClient(
        timeout=httpx.Timeout(timeout, connect=10.0), follow_redirects=True, max_redirects=5
    ) as client:
        for attempt in range(_MAX_RETRIES + 1):
            try:
                response = await client.get(url, headers=headers)
            except (httpx.TransportError, httpx.TimeoutException) as exc:
                if attempt >= _MAX_RETRIES:
                    raise
                logger.info("Transient fetch error (%s), retry %s: %s", url, attempt + 1, exc)
                await asyncio.sleep(1.5 * (2 ** attempt))  # exponential backoff (FR-CRAWL-002)
                continue
            if response.status_code in _TRANSIENT_STATUS and attempt < _MAX_RETRIES:
                logger.info("Transient %s from %s, retry %s", response.status_code, url, attempt + 1)
                await asyncio.sleep(1.5 * (2 ** attempt))  # exponential backoff (FR-CRAWL-002)
                continue
            response.raise_for_status()
            break
    declared = response.headers.get("content-length")
    if declared and declared.isdigit() and int(declared) > settings.CRAWL_MAX_BYTES:
        raise ValueError(f"Response too large ({declared} bytes)")
    if len(response.content) > settings.CRAWL_MAX_BYTES:
        raise ValueError("Response exceeds size limit")
    return response


async def fetch_raw_text(url: str, timeout: float = 20.0) -> str | None:
    """Best-effort raw GET returning the response text (for robots.txt / sitemap.xml).

    Skips the JS-render heuristic and returns None on any failure.
    """
    try:
        response = await _http_get(url, timeout)
        return response.text
    except Exception as exc:  # noqa: BLE001 - best-effort discovery
        logger.info("Raw fetch failed (%s): %s", url, exc)
        return None


def _build(url: str, title: str, text: str, content_type: str, raw_html: str | None) -> FetchedContent:
    if not text.strip():
        raise ValueError("No text content could be extracted from the URL.")
    return FetchedContent(
        url=url,
        title=title or "Untitled",
        text=text,
        content_type=content_type,
        content_hash=hashlib.sha256(text.encode("utf-8")).hexdigest(),
        content_length=len(text),
        raw_html=raw_html,
    )


async def fetch_url_content(url: str, timeout: float = 30.0) -> FetchedContent:
    """Fetch a URL and extract its text. Supports HTML/PDF/DOCX/XLSX; renders JS-only or
    blocked pages with Playwright when enabled."""
    try:
        response = await _http_get(url, timeout)
    except httpx.HTTPStatusError as exc:
        # Blocked (e.g. 403/429). Escalate: retry once with browser-like headers (some
        # sites gate non-browser UAs), then fall back to a real browser render.
        status_code = exc.response.status_code if exc.response is not None else 0
        if status_code not in (401, 403, 429):
            raise
        response = None
        try:
            response = await _http_get(url, timeout, headers=_BROWSER_HEADERS)
        except httpx.HTTPStatusError:
            response = None
        if response is None:
            if settings.CRAWL_RENDER_JS:
                rendered = await _render_with_playwright(url, timeout)
                if rendered:
                    title, text = _extract_text_from_html(rendered)
                    return _build(url, title, text, "html", rendered)
            raise

    content_type = response.headers.get("content-type", "").lower()
    lower_url = url.lower()

    if "application/pdf" in content_type or lower_url.endswith(".pdf"):
        title, text = _extract_text_from_pdf(response.content)
        return _build(url, title, text, "pdf", None)

    if "wordprocessingml" in content_type or lower_url.endswith(".docx"):
        title, text = _extract_text_from_docx(response.content)
        return _build(url, title, text, "docx", None)

    if "msword" in content_type or "rtf" in content_type or lower_url.endswith((".doc", ".rtf")):
        title, text = await _extract_text_from_doc(response.content)
        return _build(url, title, text, "doc", None)

    if "spreadsheetml" in content_type or lower_url.endswith(".xlsx"):
        title, text = _extract_text_from_xlsx(response.content)
        return _build(url, title, text, "xlsx", None)

    if "text/html" in content_type or "text/plain" in content_type or "xml" in content_type:
        raw_html = response.text
        title, text = _extract_text_from_html(raw_html)
        # Thin page (likely JS-rendered) — try a real browser render.
        if settings.CRAWL_RENDER_JS and len(text) < settings.CRAWL_THIN_TEXT_THRESHOLD:
            rendered = await _render_with_playwright(url, timeout)
            if rendered:
                r_title, r_text = _extract_text_from_html(rendered)
                if len(r_text) > len(text):
                    raw_html, title, text = rendered, (r_title or title), r_text
        return _build(url, title, text, "html", raw_html)

    raise ValueError(f"Unsupported content type: {content_type}. Only HTML, PDF, DOCX, XLSX are supported.")
